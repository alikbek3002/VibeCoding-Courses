# -*- coding: utf-8 -*-
"""Генерация редактируемого Word-документа с учебным планом Vibe Coding."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- цвета ---
TEAL = RGBColor(0x0F, 0x6E, 0x56)
TEAL_DARK = RGBColor(0x08, 0x50, 0x41)
AMBER = RGBColor(0x85, 0x4F, 0x0B)
GREY = RGBColor(0x88, 0x87, 0x80)
DARK = RGBColor(0x1A, 0x1A, 0x18)

doc = Document()

# базовый шрифт
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK


def shade(paragraph, hex_color):
    """Заливка фона абзаца."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def spacer(size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run('').font.size = Pt(size)
    return p


def title(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(2)


def subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)


def meta_line(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = True
    p.paragraph_format.space_after = Pt(12)


def month_heading(label, tag, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = color
    t = p.add_run('   ' + tag)
    t.font.size = Pt(10)
    t.bold = True
    t.font.color.rgb = color


def month_desc(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(4)


def tools(text):
    p = doc.add_paragraph()
    r = p.add_run('Инструменты: ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)


def week(num, name, sub, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run('Неделя %d — %s' % (num, name))
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = color
    p2 = doc.add_paragraph()
    r2 = p2.add_run(sub)
    r2.italic = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(4)


def session(num, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('№%d  ' % num)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.bold = bold


def callout(text, fill, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = color
    shade(p, fill)


def summary(label, items, fill, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = color
    p2 = doc.add_paragraph()
    r2 = p2.add_run(' · '.join(items))
    r2.font.size = Pt(11)
    shade(p2, fill)
    p2.paragraph_format.space_after = Pt(6)


# ============ ШАПКА ============
title('Vibe Coding — Учебный план')
subtitle('Офлайн · 3 занятия в неделю по 2 часа · 20% теория / 80% практика · Обязательное ДЗ')
meta_line('Длительность: 1–2 месяца    |    Занятий: 12 + 12    |    Старт: с нуля')

# ============ МЕСЯЦ 1 ============
month_heading('Месяц 1 — Базовый', '[ входит в курс ]', TEAL)
month_desc('Старт с нуля: учимся делать сайты словами через ИИ. Lovable, Claude и первый деплой на Vercel.')
tools('Lovable · Claude · Vercel (только в конце)')

week(1, 'Старт с Lovable', 'Занятия 1–3 · Инструмент: Lovable', TEAL_DARK)
session(1, 'Что такое ИИ и vibe coding. Регистрируемся в Lovable, пишем первый промпт — смотрим на результат')
session(2, 'Учимся описывать задачу точнее: пишем промпт, получаем лендинг, итерируем прямо в чате')
session(3, 'Делаем сайт-визитку: пишем что хотим — Lovable строит. Публикуем через встроенную кнопку')
callout('ДЗ: доработать сайт-визитку промптами и поделиться ссылкой с группой', 'E1F5EE', TEAL)

week(2, 'Переходим в Claude', 'Занятия 4–6 · Инструмент: Claude', TEAL_DARK)
session(4, 'Знакомимся с Claude: чем отличается от Lovable, как формулировать промпты, первые сайты')
session(5, 'Делаем портфолио: описываем что хотим — Claude пишет HTML и CSS. Смотрим результат в браузере')
session(6, 'Учимся уточнять и дорабатывать: «сделай кнопку зелёной», «добавь раздел с проектами» и т.д.')
callout('ДЗ: доделать портфолио промптами — сохранить файлы на компьютер', 'E1F5EE', TEAL)

week(3, 'Строим, практикуемся', 'Занятия 7–9 · Инструмент: Claude', TEAL_DARK)
session(7, 'Интерактивный сайт: форма обратной связи, кнопки, анимации — всё описываем словами в промпте')
session(8, 'Как добиться нужного дизайна промптом: цвета, шрифты, стиль — учимся формулировать красиво')
session(9, 'Мини-хакатон: каждый строит свою идею, разбираем вопросы вместе с группой')
callout('ДЗ: придумать и описать идею финального проекта — прислать преподавателю', 'E1F5EE', TEAL)

week(4, 'Финальный проект + Деплой', 'Занятия 10–12 · Claude + Vercel', TEAL_DARK)
session(10, 'Финальный проект: планируем структуру через промпт, начинаем строить — всё локально')
session(11, 'Доделываем проект. Деплой на Vercel: регистрируемся, загружаем файлы, получаем ссылку')
session(12, 'Демо-день: каждый показывает живой сайт, рассказывает про проект, получает обратную связь', bold=True)
callout('Деплой только здесь — к этому моменту проект уже готов и работает локально', 'F1EFE8', GREY)
callout('Демо-день — публичная презентация проектов', 'E1F5EE', TEAL)

summary('Что создадут ученики за первый месяц:',
        ['Лендинг', 'Сайт-визитка', 'Портфолио', 'Интерактивный сайт', 'Финальный проект'],
        'E1F5EE', TEAL)

# ============ ПЕРЕХОД ============
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Продолжение — по желанию ученика. ')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = AMBER
r2 = p.add_run('Базовый месяц — это законченный курс: ученик уже умеет делать живые сайты. '
               'Кто хочет стать про в вайбкодинге — продолжает на второй месяц: настоящие веб-приложения '
               'с базой данных, Telegram-боты и собственная CRM-система.')
r2.font.size = Pt(11)
shade(p, 'FAEEDA')

# ============ МЕСЯЦ 2 ============
month_heading('Месяц 2 — Про-уровень', '[ по желанию ]', AMBER)
month_desc('Углублённо: от сайтов к настоящим продуктам. Claude Code, красивые интерфейсы, база данных, MCP, Telegram-боты и CRM с деплоем.')
tools('Claude Code CLI · 21st.dev · Supabase · MCP серверы · Python · Vercel + Railway (только в конце)')

week(1, 'Claude Code и красивые интерфейсы', 'Занятия 1–3 · Claude Code CLI + 21st.dev', AMBER)
session(1, 'Мощный промптинг как профессионал. Claude Code CLI: устанавливаем, первый запрос в терминале — видим готовый файл')
session(2, 'Строим первый проект через Claude Code: описываем что нужно — он пишет, запускаем локально')
session(3, '21st.dev: библиотека готовых красивых компонентов. Находим нужный, вставляем промптом — доводим интерфейс')
callout('ДЗ: собрать веб-приложение с красивым интерфейсом, запустить локально', 'FAEEDA', AMBER)

week(2, 'База данных и MCP', 'Занятия 4–6 · Supabase + MCP серверы', AMBER)
session(4, 'Что такое база данных простыми словами. Supabase: создаём таблицу промптом, подключаем к проекту — данные сохраняются')
session(5, 'Добавляем авторизацию: регистрация, вход, личный кабинет — всё описываем в промпте, проверяем локально')
session(6, 'MCP серверы: подключаем GitHub MCP и Supabase MCP — Claude сам читает код и смотрит в базу. Пишем только задачу')
callout('ДЗ: приложение с авторизацией и базой данных + настроенный MCP', 'FAEEDA', AMBER)

week(3, 'Python и Telegram-боты', 'Занятия 7–9 · Python + Supabase', AMBER)
session(7, 'Python через Claude: не учим язык — просим написать, разбираем. Первый Telegram-бот через BotFather, запускаем локально')
session(8, 'Добавляем команды, кнопки и меню: пишем промптом что бот должен уметь — тестируем прямо в Telegram')
session(9, 'Бот с памятью: подключаем Supabase — бот запоминает пользователей. Объединяем команды, кнопки и базу данных')
callout('ДЗ: полноценный бот с кнопками и базой данных — работает на компьютере', 'FAEEDA', AMBER)

week(4, 'CRM-система + Деплой', 'Занятия 10–12 · Все инструменты + Vercel + Railway', AMBER)
session(10, 'Проектируем и строим CRM: интерфейс (21st.dev), база данных (Supabase), вход в систему — Claude Code собирает, всё локально')
session(11, 'Доделываем CRM: клиенты, сделки, дашборд. Деплой: CRM на Vercel, бот на Railway — живые ссылки, бот работает 24/7')
session(12, 'Демо-день: показываем CRM и бота вживую, рассказываем как строили, получаем обратную связь', bold=True)
callout('Деплой только здесь — к этому моменту всё уже работает и проверено локально', 'F1EFE8', GREY)
callout('Демо-день — публичная защита CRM-проекта', 'FAEEDA', AMBER)

summary('Что создадут ученики за второй месяц:',
        ['Full-stack веб-приложение', 'AI-агент с MCP', 'Telegram-бот с БД', 'CRM-система'],
        'FAEEDA', AMBER)

doc.save(r'C:\Users\KK\Desktop\VibeCoding-Courses\vibe_coding_plan.docx')
print('saved vibe_coding_plan.docx')
